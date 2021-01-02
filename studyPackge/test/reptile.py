# !/usr/bin/python
# -*- coding: UTF-8 -*-
# word方法库
from docx import Document

# 网页解析方法库
from bs4 import BeautifulSoup

# 网络IO包
import requests


if __name__ == '__main__':
    target = 'https://finance.sina.com.cn/fund/'
    req = requests.get(url=target)
    req.encoding = 'UTF-8'
    html = req.text
    bf = BeautifulSoup(html)
    texts = bf.find_all('ul', class_='list01')
    text = texts[3]
    aBf = BeautifulSoup(str(text))
    for link in aBf.find_all('a'):
        Doc = Document()
        name = link.get_text()
        Doc.add_heading(name)
        href = link.get('href')
        newsReq = requests.get(href)
        newsReq.encoding = 'UTF-8'
        newsHtml = newsReq.text
        newsBf = BeautifulSoup(newsHtml)
        news = newsBf.find_all('div', class_='article')
        for new in news:
            pBf = BeautifulSoup(str(new))
            ps = pBf.find_all('p')
            for p in ps:
                Doc.add_paragraph(str(p.get_text()).replace('//', '')+'\n')
        Doc.save('D:\\投资\\'+name+".docx")
        print('<<'+name+'>>' + '写入成功！')
