'''
Python CGI编程

什么是CGI
CGI 目前由NCSA维护，NCSA定义CGI如下：
CGI(Common Gateway Interface),通用网关接口,它是一段程序,运行在服务器上如：HTTP服务器，提供同客户端HTML页面的接口。

网页浏览
为了更好的了解CGI是如何工作的，我们可以从在网页上点击一个链接或URL的流程：
    1、使用你的浏览器访问URL并连接到HTTP web 服务器。
    2、Web服务器接收到请求信息后会解析URL，并查找访问的文件在服务器上是否存在，如果存在返回文件的内容，否则返回错误信息。
    3、浏览器从服务器上接收信息，并显示接收的文件或者错误信息。
CGI程序可以是Python脚本，PERL脚本，SHELL脚本，C或者C++程序等。
'''
# !/usr/bin/python
# -*- coding: UTF-8 -*-
# word方法库
from docx import Document

# 网页解析方法库
from bs4 import BeautifulSoup

# 网络IO包
import requests


if __name__ == '__main__':
    target = 'https://finance.sina.com.cn/fund/'
    req = requests.get(url=target)
    req.encoding = 'UTF-8'
    html = req.text
    bf = BeautifulSoup(html)
    texts = bf.find_all('ul', class_='list01')
    text = texts[3]
    aBf = BeautifulSoup(str(text))
    for link in aBf.find_all('a'):
        Doc = Document()
        name = link.get_text()
        Doc.add_heading(name)
        href = link.get('href')
        newsReq = requests.get(href)
        newsReq.encoding = 'UTF-8'
        newsHtml = newsReq.text
        newsBf = BeautifulSoup(newsHtml)
        news = newsBf.find_all('div', class_='article')
        for new in news:
            pBf = BeautifulSoup(str(new))
            ps = pBf.find_all('p')
            for p in ps:
                Doc.add_paragraph(str(p.get_text()).replace('//', '')+'\n')
        Doc.save('D:\\投资\\'+name+".docx")
        print('<<'+name+'>>' + '写入成功！')
